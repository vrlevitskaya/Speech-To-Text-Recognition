import requests
from docx import Document
from deep_translator import GoogleTranslator

url = "http://localhost:1337/v1/chat/completions"


def write_questions_to_docx(question, answer):
    doc = Document()
    doc.add_paragraph("Question: " + question)
    doc.add_paragraph("Answer: " + answer)
    doc.save("interactions.docx")


def translate_content(content):
    translated_content = []
    translator = GoogleTranslator(source='ru', target='en')
    for fragment in content:
        translated_fragment = translator.translate(text=fragment)
        translated_content.append(translated_fragment)
    return translated_content


def generate_answer(question, campus, dict_text):
    answer = []
    content = ['link_0']  # , 'link_1', 'link_2']
    for i in range(len(content)):
        cnt = dict_text[content[i]]
        body = {
            "model": "gpt-3.5-turbo-16k",
            "stream": False,
            'messages': [{'role': 'system', 'content': f"You are an intelligent assistant helping foreign applicants "
                                                       "in particular of HSE "
                                                       "University, answering their"
                                                       "questions about application process."
                                                       "You are obliged to Provide links and contact information. "
                                                       "You are obliged to provide information about program, "
                                                       "requirements, prices, number of paid and budget places"
                                                       "You are obliged to Limit the size of your answer to 1000 "
                                                       "characters."
                                                       "You are obliged to use information below (if it is too long, "
                                                       f"just process the first part): {cnt}"},
                         {'role': 'user', 'content': f"Answer question:{question}"
                                                     f"University {campus}."
                          }]}
        json_response = requests.post(url, json=body).json().get('choices', [])
        for choice in json_response:
            answer.append(choice.get('message', {}).get('content', ''))
    return answer
